from pathlib import Path

import pytest
from docx import Document

from app.document_service import (
    DocumentExtractionError,
    extract_docx,
    extract_text,
    extract_txt,
)


def test_extract_utf8_txt(tmp_path):
    file_path = tmp_path / "sample.txt"

    file_path.write_text(
        "AI document assistant test.",
        encoding="utf-8",
    )

    result = extract_txt(file_path)

    assert result == "AI document assistant test."


def test_extract_docx(tmp_path):
    file_path = tmp_path / "sample.docx"

    document = Document()
    document.add_paragraph(
        "First paragraph."
    )
    document.add_paragraph(
        "Second paragraph."
    )
    document.save(file_path)

    result = extract_docx(file_path)

    assert "First paragraph." in result
    assert "Second paragraph." in result


def test_extract_text_selects_txt_extractor(tmp_path):
    file_path = tmp_path / "sample.txt"

    file_path.write_text(
        "Document content.",
        encoding="utf-8",
    )

    result = extract_text(file_path)

    assert result == "Document content."


def test_unsupported_file_type_raises_error(
    tmp_path,
):
    file_path = tmp_path / "image.jpg"
    file_path.write_bytes(b"fake image")

    with pytest.raises(
        DocumentExtractionError,
        match="Unsupported document type",
    ):
        extract_text(file_path)


def test_empty_txt_raises_error(tmp_path):
    file_path = tmp_path / "empty.txt"

    file_path.write_text(
        "",
        encoding="utf-8",
    )

    with pytest.raises(
        DocumentExtractionError,
        match="No readable text",
    ):
        extract_text(file_path)